#!/usr/bin/env python3
"""
md-to-word.py v2.1.0 - Convert Markdown with Mermaid diagrams to professional Word documents

Usage:
    python md-to-word.py SOURCE.md [OUTPUT.docx]
    
Examples:
    python md-to-word.py README.md
    python md-to-word.py docs/spec.md spec.docx

Features:
    - Mermaid diagrams → PNG (90% page fit, preserves aspect ratio)
    - Table formatting: Microsoft blue headers, borders, alternating rows
    - Table pagination: prevents orphan headers, keeps rows together
    - SVG banners → PNG for Word compatibility
    - Markdown preprocessing: fixes bullet lists, checkboxes

Requirements:
    - pandoc (winget install pandoc)
    - mermaid-cli (npm install -g @mermaid-js/mermaid-cli)
    - python-docx (pip install python-docx)
    - svgexport (npm install -g svgexport) [optional, for SVG banners]
"""

import argparse
import os
import re
import struct
import subprocess
import sys
import tempfile
from pathlib import Path

# Check for python-docx
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import nsdecls, qn
    from docx.oxml import parse_xml
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

# Page Layout Constants (Letter: 8.5" x 11", 1" margins)
PAGE_WIDTH_INCHES = 6.5   # Usable width (8.5" - 1" margins each side)
PAGE_HEIGHT_INCHES = 9.0  # Usable height (11" - 1" margins each side)
MAX_IMAGE_RATIO = 0.90    # 90% of page max
MAX_IMAGE_WIDTH = PAGE_WIDTH_INCHES * MAX_IMAGE_RATIO   # ~5.85"
MAX_IMAGE_HEIGHT = PAGE_HEIGHT_INCHES * MAX_IMAGE_RATIO # ~8.1"
PNG_DPI = 96  # Standard screen DPI for mermaid-cli output


def get_png_dimensions(png_path: Path) -> tuple[int, int]:
    """Read PNG dimensions from file header (no Pillow dependency).
    
    Returns (width, height) in pixels, or (0, 0) if failed.
    """
    try:
        with open(png_path, 'rb') as f:
            # PNG signature (8 bytes) + IHDR chunk length (4 bytes) + IHDR type (4 bytes)
            header = f.read(24)
            if header[:8] == b'\x89PNG\r\n\x1a\n':
                # Width and height are at bytes 16-20 and 20-24 (big-endian)
                width = struct.unpack('>I', header[16:20])[0]
                height = struct.unpack('>I', header[20:24])[0]
                return width, height
    except (IOError, struct.error):
        pass
    return 0, 0


def calculate_optimal_size(png_path: Path, mmd_content: str) -> str:
    """Calculate optimal image size to fit 90% of page both horizontally and vertically.
    
    Reads actual PNG dimensions and scales to fit within:
    - Max width: 5.85" (90% of 6.5" usable width)
    - Max height: 8.1" (90% of 9" usable height)
    
    Only specifies ONE dimension (width OR height) - pandoc preserves aspect ratio.
    """
    width_px, height_px = get_png_dimensions(png_path)
    
    if width_px == 0 or height_px == 0:
        # Fallback to heuristic-based sizing if PNG read fails
        return determine_image_size_heuristic(mmd_content)
    
    # Convert pixels to inches (mermaid-cli uses 96 DPI)
    width_in = width_px / PNG_DPI
    height_in = height_px / PNG_DPI
    
    # Calculate scale factors to fit within 90% bounds
    width_scale = MAX_IMAGE_WIDTH / width_in if width_in > 0 else 1
    height_scale = MAX_IMAGE_HEIGHT / height_in if height_in > 0 else 1
    
    # Use the smaller scale factor (most constraining)
    scale = min(width_scale, height_scale, 1.0)  # Never upscale
    
    # Calculate target dimensions
    target_width = width_in * scale
    target_height = height_in * scale
    
    # Determine which dimension to specify based on aspect ratio
    # For wide images, specify width; for tall images, specify height
    aspect_ratio = width_in / height_in if height_in > 0 else 1
    
    if aspect_ratio >= 1.0:
        # Wide or square: constrain by width
        return f'{{width={target_width:.1f}in}}'
    else:
        # Tall: constrain by height (ensures it fits on page)
        return f'{{height={target_height:.1f}in}}'


def determine_image_size_heuristic(mmd_content: str) -> str:
    """Fallback heuristic sizing when PNG dimensions unavailable."""
    content_lower = mmd_content.lower()
    subgraph_count = len(re.findall(r'subgraph', content_lower))
    
    if 'gantt' in content_lower:
        return '{width=5.8in}'
    if subgraph_count >= 3:
        return '{width=5.8in}'
    if 'flowchart lr' in content_lower or 'graph lr' in content_lower:
        return '{width=5.8in}'
    if 'flowchart tb' in content_lower or 'graph tb' in content_lower:
        if subgraph_count >= 2:
            return '{height=8in}'  # Tall diagram - use height constraint
        return '{width=5in}'
    return '{width=5in}'


def find_mermaid_blocks(content: str) -> list[tuple[int, str]]:
    """Find all mermaid code blocks and return (index, content) tuples."""
    pattern = r'```mermaid\r?\n(.*?)```'
    matches = list(re.finditer(pattern, content, re.DOTALL))
    return [(i, m.group(1)) for i, m in enumerate(matches)]


def convert_mermaid_to_png(mmd_content: str, output_path: Path) -> bool:
    """Convert mermaid content to PNG using mmdc."""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False, encoding='utf-8') as f:
        f.write(mmd_content)
        temp_mmd = f.name
    
    try:
        result = subprocess.run(
            f'npx mmdc -i "{temp_mmd}" -o "{output_path}" -b white',
            capture_output=True,
            text=True,
            shell=True
        )
        return result.returncode == 0
    finally:
        os.unlink(temp_mmd)


def convert_svg_to_png(svg_path: Path, png_path: Path) -> bool:
    """Convert SVG to PNG using svgexport."""
    try:
        result = subprocess.run(
            f'npx svgexport "{svg_path}" "{png_path}" 800:',
            capture_output=True,
            text=True,
            shell=True
        )
        return result.returncode == 0
    except FileNotFoundError:
        print(f"WARNING: svgexport not available, skipping {svg_path}")
        return False


def set_cell_shading(cell, color: str):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def autofit_table(table):
    """Set table to auto-fit to content."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(
            '<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        tbl.insert(0, tblPr)
    
    # Remove existing width settings
    for existing in tblPr.findall(qn('w:tblW')):
        tblPr.remove(existing)
    
    # Set auto width (fit to content)
    tbl_width = parse_xml(
        '<w:tblW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:type="auto" w:w="0"/>'
    )
    tblPr.append(tbl_width)
    
    # Set table layout to auto (allows columns to resize)
    for existing in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(existing)
    layout = parse_xml(
        '<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:type="autofit"/>'
    )
    tblPr.append(layout)


def set_table_borders(table):
    """Apply professional borders to table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(
            '<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        tbl.insert(0, tblPr)
    
    borders = parse_xml(
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="6" w:space="0" w:color="666666"/>'
        '<w:left w:val="single" w:sz="6" w:space="0" w:color="666666"/>'
        '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="666666"/>'
        '<w:right w:val="single" w:sz="6" w:space="0" w:color="666666"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        '</w:tblBorders>'
    )
    
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    tblPr.append(borders)


def set_row_cant_split(row):
    """Prevent a table row from splitting across pages."""
    trPr = row._tr.get_or_add_trPr()
    # Remove existing cantSplit if present
    for existing in trPr.findall(qn('w:cantSplit')):
        trPr.remove(existing)
    cant_split = parse_xml(
        '<w:cantSplit xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )
    trPr.append(cant_split)


def set_table_keep_together(table, max_rows_to_keep=12):
    """Keep small tables together on one page by setting keepNext on rows.
    
    For tables with <= max_rows_to_keep rows, keeps entire table together.
    For larger tables, at minimum keeps header + first data row together.
    """
    num_rows = len(table.rows)
    
    # Determine how many rows to keep together
    if num_rows <= max_rows_to_keep:
        # Small table - keep all rows together
        rows_to_link = num_rows - 1  # All but the last row
    else:
        # Large table - keep header + at least 2 data rows together
        rows_to_link = min(3, num_rows - 1)
    
    for i, row in enumerate(table.rows):
        # Prevent individual rows from splitting across pages
        set_row_cant_split(row)
        
        # Set keepNext on rows to keep table together
        if i < rows_to_link:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    pf = paragraph.paragraph_format
                    pf.keep_with_next = True
                    pf.keep_together = True


def format_tables(doc: Document):
    """Apply professional formatting to all tables."""
    for table in doc.tables:
        autofit_table(table)
        set_table_borders(table)
        set_table_keep_together(table)  # Smart pagination
        
        # Header row - Microsoft blue with white text
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                set_cell_shading(cell, '0078D4')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(10)
        
        # Data rows - alternating colors
        for i, row in enumerate(table.rows[1:], 1):
            for cell in row.cells:
                color = 'F0F0F0' if i % 2 == 0 else 'FFFFFF'
                set_cell_shading(cell, color)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 0, 0)


def center_images(doc: Document):
    """Center all images (inline shapes) in their paragraphs."""
    for paragraph in doc.paragraphs:
        # Check if paragraph contains an image
        inline_shapes = paragraph._element.findall('.//wp:inline', 
            {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
        if inline_shapes:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def fix_paragraph_spacing(doc: Document):
    """Ensure proper spacing and widow/orphan control for all paragraphs."""
    for paragraph in doc.paragraphs:
        # Skip empty paragraphs
        if not paragraph.text.strip():
            continue
        
        pf = paragraph.paragraph_format
        
        # Enable widow/orphan control (prevents single lines at page top/bottom)
        pf.widow_control = True
        
        # List items should have less spacing
        if paragraph.style.name and 'List' in paragraph.style.name:
            pf.space_before = Pt(2)
            pf.space_after = Pt(2)
        # Regular paragraphs
        elif paragraph.style.name == 'Normal':
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)


def format_headings(doc: Document):
    """Apply consistent heading styles with orphan prevention."""
    heading_colors = {
        'Heading 1': RGBColor(0x00, 0x52, 0x8B),  # Dark blue
        'Heading 2': RGBColor(0x00, 0x78, 0xD4),  # Microsoft blue
        'Heading 3': RGBColor(0x10, 0x5E, 0x7E),  # Teal
        'Heading 4': RGBColor(0x10, 0x5E, 0x7E),  # Teal
    }
    
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else None
        
        if style_name and style_name.startswith('Heading'):
            # Apply colors
            if style_name in heading_colors:
                for run in paragraph.runs:
                    run.font.color.rgb = heading_colors[style_name]
            
            # Paragraph formatting
            pf = paragraph.paragraph_format
            
            # CRITICAL: Keep heading with next paragraph (prevents orphan titles)
            pf.keep_with_next = True
            
            # Also prevent page breaks within the heading itself
            pf.keep_together = True
            
            # Spacing based on heading level
            if style_name == 'Heading 1':
                pf.space_before = Pt(18)
                pf.space_after = Pt(6)
            elif style_name == 'Heading 2':
                pf.space_before = Pt(14)
                pf.space_after = Pt(4)
            elif style_name == 'Heading 3':
                pf.space_before = Pt(12)
                pf.space_after = Pt(4)
            elif style_name == 'Heading 4':
                pf.space_before = Pt(10)
                pf.space_after = Pt(3)


def format_code_blocks(doc: Document):
    """Apply professional formatting to code blocks.
    
    Pandoc generates code blocks with styles like 'Source Code', 'Verbatim Char',
    or paragraphs containing code runs. This function:
    - Applies Consolas monospace font
    - Sets light gray background
    - Adds subtle border
    - Uses smaller font size
    """
    # Code-related style names that pandoc may generate
    code_styles = {'Source Code', 'Verbatim Char', 'Code', 'SourceCode'}
    
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ''
        
        # Check if this is a code block paragraph
        is_code = (
            style_name in code_styles or
            'code' in style_name.lower() or
            'verbatim' in style_name.lower()
        )
        
        if is_code:
            # Format the paragraph
            pf = paragraph.paragraph_format
            pf.space_before = Pt(3)
            pf.space_after = Pt(3)
            pf.keep_together = True  # Don't split code blocks
            
            # Format each run with monospace font
            for run in paragraph.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)  # Dark gray text
            
            # Apply shading to the paragraph (light gray background)
            shading_elm = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>'
            )
            paragraph._element.get_or_add_pPr().append(shading_elm)
            
            # Add left border for code block visual distinction
            pBdr = parse_xml(
                '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:left w:val="single" w:sz="24" w:space="4" w:color="CCCCCC"/>'
                '<w:top w:val="single" w:sz="4" w:space="1" w:color="E0E0E0"/>'
                '<w:bottom w:val="single" w:sz="4" w:space="1" w:color="E0E0E0"/>'
                '<w:right w:val="single" w:sz="4" w:space="1" w:color="E0E0E0"/>'
                '</w:pBdr>'
            )
            pPr = paragraph._element.get_or_add_pPr()
            # Remove existing borders
            for existing in pPr.findall(qn('w:pBdr')):
                pPr.remove(existing)
            pPr.append(pBdr)


def apply_all_formatting(doc: Document, format_tables_flag: bool = True):
    """Apply all formatting improvements to the document."""
    if format_tables_flag:
        format_tables(doc)
    center_images(doc)
    format_headings(doc)
    format_code_blocks(doc)
    fix_paragraph_spacing(doc)


def preprocess_markdown(content: str) -> str:
    """
    Fix common markdown issues before pandoc conversion.
    
    Issues addressed:
    - Bullet lists running together (need blank lines)
    - Inconsistent list markers
    - Missing blank lines after headings
    - Checkbox lists
    - Blockquote formatting
    """
    lines = content.split('\n')
    result = []
    prev_was_list = False
    prev_was_blank = False
    
    for i, line in enumerate(lines):
        stripped = line.strip()
        
        # Detect list items (-, *, numbered, checkbox)
        is_list = bool(re.match(r'^[-*+]\s|^\d+\.\s|^[-*+]\s*\[[ xX]\]', stripped))
        is_blank = not stripped
        is_heading = stripped.startswith('#')
        is_blockquote = stripped.startswith('>')
        is_table_row = '|' in stripped and stripped.startswith('|')
        is_code_fence = stripped.startswith('```')
        
        # Add blank line before lists if previous line was not blank/list
        if is_list and not prev_was_list and not prev_was_blank and result:
            result.append('')
        
        # Add blank line after heading if next line is not blank
        if i > 0 and result:
            prev_line = lines[i-1].strip()
            if prev_line.startswith('#') and not is_blank:
                # Check if we didn't already add a blank
                if result[-1] != '':
                    result.append('')
        
        # Convert checkbox markers for pandoc compatibility
        # - [ ] item -> - ☐ item
        # - [x] item -> - ☑ item
        if re.match(r'^[-*+]\s*\[[ ]\]', stripped):
            line = re.sub(r'^([-*+])\s*\[ \]', r'\1 ☐', line)
        elif re.match(r'^[-*+]\s*\[[xX]\]', stripped):
            line = re.sub(r'^([-*+])\s*\[[xX]\]', r'\1 ☑', line)
        
        result.append(line)
        prev_was_list = is_list
        prev_was_blank = is_blank
    
    # Add blank line after lists before non-list content
    final = []
    for i, line in enumerate(result):
        final.append(line)
        stripped = line.strip()
        is_list = bool(re.match(r'^[-*+]\s|^\d+\.\s|^[-*+]\s*[☐☑]', stripped))
        
        if is_list and i + 1 < len(result):
            next_stripped = result[i + 1].strip()
            next_is_list = bool(re.match(r'^[-*+]\s|^\d+\.\s|^[-*+]\s*[☐☑]', next_stripped))
            next_is_blank = not next_stripped
            
            # If next line is not a list item and not blank, add blank
            if not next_is_list and not next_is_blank:
                final.append('')
    
    return '\n'.join(final)


def main():
    parser = argparse.ArgumentParser(
        description='Convert Markdown with Mermaid diagrams to Word document'
    )
    parser.add_argument('source', help='Source Markdown file')
    parser.add_argument('output', nargs='?', help='Output Word file (default: source.docx)')
    parser.add_argument('--images-dir', default='images', help='Directory for generated images')
    parser.add_argument('--no-format-tables', action='store_true', help='Skip table formatting')
    parser.add_argument('--keep-temp', action='store_true', help='Keep temporary files')
    
    args = parser.parse_args()
    
    source_path = Path(args.source)
    if not source_path.exists():
        print(f"ERROR: Source file not found: {source_path}")
        sys.exit(1)
    
    output_path = Path(args.output) if args.output else source_path.with_suffix('.docx')
    images_dir = source_path.parent / args.images_dir
    images_dir.mkdir(exist_ok=True)
    
    print(f"📄 Converting {source_path} → {output_path}")
    
    # Read source
    content = source_path.read_text(encoding='utf-8')
    
    # Phase 0: Preprocess markdown to fix formatting issues
    print(f"🔧 Preprocessing markdown...")
    content = preprocess_markdown(content)
    
    # Phase 1: Find and convert Mermaid diagrams
    mermaid_blocks = find_mermaid_blocks(content)
    print(f"📊 Found {len(mermaid_blocks)} Mermaid diagrams")
    
    replacements = []
    for idx, mmd_content in mermaid_blocks:
        png_name = f'diagram-{idx + 1}.png'
        png_path = images_dir / png_name
        
        print(f"   Converting diagram {idx + 1}...", end=' ')
        if convert_mermaid_to_png(mmd_content, png_path):
            # Calculate optimal size from actual PNG dimensions
            size = calculate_optimal_size(png_path, mmd_content)
            replacements.append(f'![Diagram {idx + 1}]({args.images_dir}/{png_name}){size}')
            print(f"✓ {size}")
        else:
            print("✗ (failed)")
            replacements.append(f'![Diagram {idx + 1}]({args.images_dir}/{png_name})')
    
    # Phase 2: Convert SVG references to PNG
    svg_pattern = r'!\[([^\]]*)\]\(([^)]+\.svg)\)'
    for match in re.finditer(svg_pattern, content):
        alt_text, svg_rel_path = match.groups()
        svg_path = source_path.parent / svg_rel_path
        
        if svg_path.exists():
            png_name = svg_path.stem + '.png'
            png_path = images_dir / png_name
            
            if not png_path.exists():
                print(f"🖼️  Converting SVG: {svg_path.name}...", end=' ')
                if convert_svg_to_png(svg_path, png_path):
                    print("✓")
                else:
                    print("✗")
            
            # Update content to use PNG (90% max width = 5.8in)
            new_ref = f'![{alt_text}]({args.images_dir}/{png_name}){{width=5.8in}}'
            content = content.replace(match.group(0), new_ref)
    
    # Phase 3: Replace mermaid blocks with image references
    pattern = r'```mermaid\r?\n.*?```'
    for replacement in replacements:
        content = re.sub(pattern, replacement, content, count=1, flags=re.DOTALL)
    
    # Write temporary markdown
    temp_md = source_path.parent / '_temp_word.md'
    temp_md.write_text(content, encoding='utf-8')
    
    # Phase 4: Convert to Word with pandoc
    # Use --resource-path so pandoc resolves relative image paths from source dir 
    print(f"📝 Generating Word document...")
    resource_path = source_path.parent.resolve()
    result = subprocess.run(
        f'pandoc "{temp_md}" -o "{output_path}" --from markdown --to docx --resource-path="{resource_path}"',
        capture_output=True,
        text=True,
        shell=True
    )
    
    if result.returncode != 0:
        print(f"ERROR: pandoc failed: {result.stderr}")
        sys.exit(1)
    
    # Phase 5: Apply all formatting (tables, images, headings, spacing)
    print(f"🎨 Applying formatting...")
    doc = Document(str(output_path))
    apply_all_formatting(doc, format_tables_flag=not args.no_format_tables)
    doc.save(str(output_path))
    
    # Cleanup
    if not args.keep_temp:
        temp_md.unlink()
    
    print(f"✅ Done! Output: {output_path}")
    print(f"   Size: {output_path.stat().st_size / 1024:.1f} KB")


if __name__ == '__main__':
    main()
